"""Render a meeting protocol as PDF or DOCX.

The PDF is a Jinja-rendered HTML run through WeasyPrint; the DOCX is
built directly via python-docx (different visual treatment — Word users
expect a clean structured document, not a styled marketing pdf).

Both formats share the same context-building stage that resolves
{{speaker:N}} tokens, parses markdown sections out of meeting.protocol,
and pulls participant + task + author data.
"""
from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable

from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.meeting import Meeting
from app.models.meeting_task import MeetingTask
from app.models.participant import MeetingSpeaker, Participant
from app.models.profile import Profile
from app.models.user import User


_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
_jinja = Environment(
    loader=FileSystemLoader(str(_TEMPLATES_DIR)),
    autoescape=select_autoescape(["html"]),
    trim_blocks=False,
    lstrip_blocks=False,
)

_AVATAR_COLORS = ("violet", "cyan", "pink", "amber", "emerald")
_RU_MONTHS = (
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
)
_SPEAKER_TOKEN_RE = re.compile(r"\{\{speaker:(\d+)\}\}")


# ─────────────────────────────────────────────────────────────────────
# Data shapes (kept as plain dicts inside the Jinja context to avoid
# coupling the template to dataclass shape — but services use these to
# stay readable)
# ─────────────────────────────────────────────────────────────────────


@dataclass
class _Author:
    name: str
    position: str | None = None
    company: str | None = None


@dataclass
class _Topic:
    title: str
    bullets: list[str] = field(default_factory=list)


@dataclass
class _TaskRow:
    title: str
    description: str | None
    owner: str | None
    due: str | None


# ─────────────────────────────────────────────────────────────────────
# Loaders
# ─────────────────────────────────────────────────────────────────────


async def _load_meeting(
    db: AsyncSession, meeting_id: uuid.UUID, user_id: uuid.UUID
) -> Meeting:
    q = await db.execute(
        select(Meeting).where(Meeting.id == meeting_id, Meeting.user_id == user_id)
    )
    meeting = q.scalar_one_or_none()
    if meeting is None:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Meeting not found")
    return meeting


async def _load_speakers_with_participants(
    db: AsyncSession, meeting_id: uuid.UUID
) -> list[tuple[MeetingSpeaker, Participant | None]]:
    q = await db.execute(
        select(MeetingSpeaker, Participant)
        .outerjoin(Participant, Participant.id == MeetingSpeaker.participant_id)
        .where(MeetingSpeaker.meeting_id == meeting_id)
        .order_by(MeetingSpeaker.speaker_label.asc())
    )
    return list(q.all())


async def _load_tasks(
    db: AsyncSession, meeting_id: uuid.UUID
) -> list[tuple[MeetingTask, Participant | None]]:
    q = await db.execute(
        select(MeetingTask, Participant)
        .outerjoin(Participant, Participant.id == MeetingTask.assignee_participant_id)
        .where(MeetingTask.meeting_id == meeting_id)
        .order_by(MeetingTask.position.asc(), MeetingTask.created_at.asc())
    )
    return list(q.all())


async def _load_author(db: AsyncSession, user_id: uuid.UUID) -> _Author:
    q = await db.execute(
        select(Profile, User)
        .join(User, User.id == Profile.user_id)
        .where(Profile.user_id == user_id)
    )
    row = q.one_or_none()
    if row is None:
        # Fallback to email-only
        u_q = await db.execute(select(User).where(User.id == user_id))
        u = u_q.scalar_one_or_none()
        return _Author(name=(u.email if u else "—"))
    profile, user = row
    name = (profile.full_name or "").strip() or user.email or "—"
    return _Author(
        name=name,
        position=(profile.position or "").strip() or None,
        company=(profile.company_name or "").strip() or None,
    )


# ─────────────────────────────────────────────────────────────────────
# Markdown protocol parsing — extract ## sections + flatten bullets
# ─────────────────────────────────────────────────────────────────────


def _resolve_tokens(text: str, label_to_name: dict[str, str]) -> str:
    def _sub(m: re.Match[str]) -> str:
        n = m.group(1)
        return label_to_name.get(f"SPEAKER_{n}", f"Спикер {int(n) + 1}")

    return _SPEAKER_TOKEN_RE.sub(_sub, text)


def _split_sections(md: str) -> dict[str, str]:
    """Splits markdown into {section_name_lowercased: section_body}.
    Section name = text after `## `. Body = lines until the next `## `.
    Anything before the first `## ` (e.g. the `# Title`) is dropped.
    """
    out: dict[str, str] = {}
    current: str | None = None
    buf: list[str] = []
    for line in md.splitlines():
        stripped = line.lstrip()
        if stripped.startswith("## "):
            if current is not None:
                out[current] = "\n".join(buf).strip()
            current = stripped[3:].strip().lower()
            buf = []
        elif stripped.startswith("# "):
            # main title — skip; we use meeting.title instead
            continue
        else:
            if current is not None:
                buf.append(line)
    if current is not None:
        out[current] = "\n".join(buf).strip()
    return out


def _bullets(section_md: str) -> list[str]:
    """Extracts top-level `- ` bullets from a section body. Continuation
    lines (indented or non-empty without bullet marker) are appended to
    the current bullet."""
    items: list[str] = []
    cur: list[str] = []
    for line in section_md.splitlines():
        if re.match(r"^\s*[-*]\s+", line):
            if cur:
                items.append(" ".join(cur).strip())
                cur = []
            cur.append(re.sub(r"^\s*[-*]\s+", "", line).strip())
        elif line.strip() == "":
            if cur:
                items.append(" ".join(cur).strip())
                cur = []
        else:
            if cur:
                cur.append(line.strip())
    if cur:
        items.append(" ".join(cur).strip())
    return [it for it in items if it]


# ─────────────────────────────────────────────────────────────────────
# Context builder
# ─────────────────────────────────────────────────────────────────────


def _ru_date(d: datetime) -> str:
    return f"{d.day} {_RU_MONTHS[d.month - 1]} {d.year}, {d.hour:02d}:{d.minute:02d}"


def _ru_short_date(d: datetime) -> str:
    return f"{d.day} {_RU_MONTHS[d.month - 1]}"


def _ru_duration(seconds: int | None) -> str:
    if not seconds or seconds <= 0:
        return "—"
    minutes = seconds // 60
    if minutes < 60:
        return f"{minutes} мин"
    hours = minutes // 60
    rem = minutes % 60
    if rem == 0:
        return f"{hours} ч"
    return f"{hours} ч {rem} мин"


def _initial(name: str) -> str:
    parts = [p for p in name.strip().split() if p]
    if not parts:
        return "?"
    return parts[0][0].upper()


def _doc_id(meeting_id: uuid.UUID, created_at: datetime) -> str:
    suffix = meeting_id.hex[:3].upper()
    return f"BRF-{created_at.year}-{created_at.month:02d}{created_at.day:02d}-{suffix}"


async def _build_context(
    db: AsyncSession, meeting: Meeting
) -> dict:
    speaker_rows = await _load_speakers_with_participants(db, meeting.id)
    task_rows = await _load_tasks(db, meeting.id)
    author = await _load_author(db, meeting.user_id)

    # SPEAKER_N → display name
    label_to_name: dict[str, str] = {}
    participants: list[dict] = []
    seen_pids: set[uuid.UUID] = set()
    for idx, (sp, p) in enumerate(speaker_rows):
        if p is not None:
            label_to_name[sp.speaker_label] = p.name
            if p.id not in seen_pids:
                seen_pids.add(p.id)
                participants.append({
                    "name": p.name,
                    "role": (p.role or "").strip() or None,
                    "color": _AVATAR_COLORS[len(participants) % len(_AVATAR_COLORS)],
                    "initial": _initial(p.name),
                })
        else:
            # Unbound speaker — show as "Спикер N" placeholder
            try:
                n = int(sp.speaker_label.split("_")[-1])
            except ValueError:
                n = idx
            label_to_name[sp.speaker_label] = f"Спикер {n + 1}"

    md = (meeting.protocol or "")
    md_resolved = _resolve_tokens(md, label_to_name)
    sections = _split_sections(md_resolved)

    # Aliases — accept several possible heading spellings
    def _section(*names: str) -> str:
        for n in names:
            if n in sections and sections[n]:
                return sections[n]
        return ""

    summary = _section("резюме", "summary", "сводка", "краткое содержание")
    topics_md = _section("темы", "ключевые темы", "topics")
    decisions_md = _section("решения", "принятые решения", "decisions")

    topics = [{"title": t, "bullets": []} for t in _bullets(topics_md)] if topics_md else []
    decisions = _bullets(decisions_md) if decisions_md else []

    # Tasks (structured table from MeetingTask)
    tasks: list[dict] = []
    for t, owner in task_rows:
        if t.status == "cancelled":
            continue  # skip cancelled tasks from the deliverable
        tasks.append({
            "title": _resolve_tokens(t.title or "", label_to_name),
            "description": _resolve_tokens(t.description or "", label_to_name) or None,
            "owner": owner.name if owner is not None else None,
            "due": _ru_short_date(t.due_date) if t.due_date else None,
        })

    has_page_2 = bool(tasks or decisions or author)

    return {
        "doc_id": _doc_id(meeting.id, meeting.created_at),
        "generated_at": _ru_date(datetime.now()),
        "title": (meeting.title or "Встреча").strip(),
        "eyebrow": None,  # could populate from series in future
        "meta": {
            "date": _ru_date(meeting.created_at),
            "duration": _ru_duration(meeting.duration_seconds),
            "participants_count": f"{len(participants)} {'спикер' if len(participants) == 1 else 'спикеров'}",
            "tasks_count": str(len(tasks)),
        },
        "participants": participants,
        "summary": summary or None,
        "topics": topics,
        "decisions": decisions,
        "tasks": tasks,
        "author": {
            "name": author.name,
            "position": author.position,
            "company": author.company,
        } if author and author.name else None,
        "has_page_2": has_page_2,
        "total_pages": 2 if has_page_2 else 1,
    }


# ─────────────────────────────────────────────────────────────────────
# PDF rendering — WeasyPrint
# ─────────────────────────────────────────────────────────────────────


async def render_pdf(
    db: AsyncSession, meeting_id: uuid.UUID, user_id: uuid.UUID
) -> tuple[bytes, str]:
    """Returns (pdf_bytes, suggested_filename)."""
    # Imported lazily so the rest of the app boots even when WeasyPrint's
    # native deps (cairo, pango) are missing — only the export endpoint
    # needs them.
    from weasyprint import HTML  # type: ignore

    meeting = await _load_meeting(db, meeting_id, user_id)
    ctx = await _build_context(db, meeting)
    template = _jinja.get_template("protocol_pdf.html")
    html_str = template.render(**ctx)
    pdf_bytes = HTML(string=html_str).write_pdf()
    filename = _safe_filename(meeting.title, meeting.created_at, "pdf")
    return pdf_bytes, filename


# ─────────────────────────────────────────────────────────────────────
# DOCX rendering — python-docx
# ─────────────────────────────────────────────────────────────────────


async def render_docx(
    db: AsyncSession, meeting_id: uuid.UUID, user_id: uuid.UUID
) -> tuple[bytes, str]:
    """Build a clean structured Word document. No gradients, no rounded
    boxes — Word users expect a normal looking doc, not a marketing pdf.
    Sections: Title, meta paragraph, Участники, Резюме, Темы, Решения,
    Задачи (table), author footer.
    """
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor

    meeting = await _load_meeting(db, meeting_id, user_id)
    ctx = await _build_context(db, meeting)

    doc = Document()

    # Title
    title_p = doc.add_heading(ctx["title"], level=0)
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x0B, 0x33)

    # Meta line
    meta_line = (
        f"{ctx['meta']['date']} · {ctx['meta']['duration']} · "
        f"{ctx['meta']['participants_count']} · {ctx['doc_id']}"
    )
    p = doc.add_paragraph()
    run = p.add_run(meta_line)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x60, 0x80)

    # Participants
    if ctx["participants"]:
        doc.add_heading("Участники", level=2)
        for prt in ctx["participants"]:
            line = prt["name"]
            if prt.get("role"):
                line += f" — {prt['role']}"
            doc.add_paragraph(line, style="List Bullet")

    # Summary
    if ctx["summary"]:
        doc.add_heading("Резюме", level=2)
        doc.add_paragraph(ctx["summary"])

    # Topics
    if ctx["topics"]:
        doc.add_heading("Ключевые темы", level=2)
        for i, t in enumerate(ctx["topics"], start=1):
            p = doc.add_paragraph()
            run_n = p.add_run(f"{i}. ")
            run_n.bold = True
            p.add_run(t["title"])
            for b in t.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet 2")

    # Decisions
    if ctx["decisions"]:
        doc.add_heading("Принятые решения", level=2)
        for d in ctx["decisions"]:
            doc.add_paragraph(d, style="List Bullet")

    # Tasks
    if ctx["tasks"]:
        doc.add_heading("Задачи", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "№"
        hdr[1].text = "Задача"
        hdr[2].text = "Ответственный"
        hdr[3].text = "Срок"
        for i, t in enumerate(ctx["tasks"], start=1):
            row = table.add_row().cells
            row[0].text = f"{i:02d}"
            row[1].text = t["title"]
            if t.get("description"):
                row[1].add_paragraph(t["description"]).runs[0].italic = True
            row[2].text = t.get("owner") or "—"
            row[3].text = t.get("due") or "—"

    # Author footer
    if ctx["author"]:
        doc.add_paragraph()
        a = ctx["author"]
        line = f"Протокол подготовил(а): {a['name']}"
        if a.get("position"):
            line += f", {a['position']}"
        if a.get("company"):
            line += f" · {a['company']}"
        line += "."
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6B, 0x60, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    filename = _safe_filename(meeting.title, meeting.created_at, "docx")
    return buf.getvalue(), filename


# ─────────────────────────────────────────────────────────────────────
# Filename helper — strips characters that break common file systems
# and email clients.
# ─────────────────────────────────────────────────────────────────────


def _safe_filename(title: str | None, created_at: datetime, ext: str) -> str:
    base = (title or "Протокол").strip()
    base = re.sub(r"[\\/:*?\"<>|\r\n\t]+", " ", base)
    base = re.sub(r"\s+", " ", base).strip()[:80] or "Протокол"
    stamp = created_at.strftime("%Y-%m-%d")
    return f"{base} · {stamp}.{ext}"


async def render_for_format(
    db: AsyncSession, meeting_id: uuid.UUID, user_id: uuid.UUID, fmt: str
) -> tuple[bytes, str, str]:
    """Returns (bytes, filename, content_type)."""
    if fmt == "pdf":
        data, name = await render_pdf(db, meeting_id, user_id)
        return data, name, "application/pdf"
    if fmt == "docx":
        data, name = await render_docx(db, meeting_id, user_id)
        return data, name, (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    from fastapi import HTTPException
    raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")
